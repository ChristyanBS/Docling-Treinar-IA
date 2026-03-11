"""
Serviço de processamento e treinamento com documentos.
Armazena tudo no PostgreSQL (documents + knowledge_chunks) com embeddings pgvector.
"""

import os
import re
import json
import logging
from typing import List
from datetime import datetime

from sqlalchemy import text as sa_text

from app.config import UPLOAD_DIR, CHUNKS_DIR
from app.database import SessionLocal, Document, KnowledgeChunk, TrainingStatus
from app.embedding_service import embed_texts, embed_single

logger = logging.getLogger(__name__)

# Distância cosseno máxima para considerar um chunk relevante (0=idêntico, 2=oposto)
# Chunks com distância > 0.85 são filtrados como irrelevantes
MAX_COSINE_DISTANCE = 0.80


def split_text_into_chunks(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
    """Divide texto em chunks semânticos: primeiro por seções/parágrafos, depois por tamanho."""
    if not text or not text.strip():
        return []

    # Dividir por seções markdown (## Título) ou blocos de linhas em branco duplas
    sections = re.split(r'\n(?=#{1,3}\s)|\n{3,}', text)

    chunks = []
    for section in sections:
        section = section.strip()
        if not section:
            continue

        words = section.split()
        if len(words) <= chunk_size:
            # Seção cabe inteira num chunk
            if len(words) >= 10:  # mínimo 10 palavras para ser útil
                chunks.append(section)
        else:
            # Seção grande: dividir por parágrafos primeiro
            paragraphs = re.split(r'\n\n+', section)
            current_chunk = []
            current_size = 0

            for para in paragraphs:
                para_words = len(para.split())
                if current_size + para_words > chunk_size and current_chunk:
                    chunks.append("\n\n".join(current_chunk).strip())
                    # Overlap: manter último parágrafo
                    if overlap > 0 and current_chunk:
                        last = current_chunk[-1]
                        current_chunk = [last]
                        current_size = len(last.split())
                    else:
                        current_chunk = []
                        current_size = 0
                current_chunk.append(para)
                current_size += para_words

            if current_chunk:
                text_chunk = "\n\n".join(current_chunk).strip()
                if len(text_chunk.split()) >= 10:
                    chunks.append(text_chunk)

    # Fallback: se não gerou nada, usar divisão mecânica
    if not chunks:
        words = text.split()
        start = 0
        while start < len(words):
            end = start + chunk_size
            chunk = " ".join(words[start:end])
            if chunk.strip():
                chunks.append(chunk.strip())
            start += chunk_size - overlap

    return chunks


def extract_text_from_file(filepath: str) -> str:
    """Extrai texto de diversos formatos de arquivo."""
    ext = os.path.splitext(filepath)[1].lower()
    try:
        # Docling (principal)
        if ext in (".pdf", ".docx", ".doc", ".pptx", ".html", ".md"):
            try:
                from docling.document_converter import DocumentConverter
                converter = DocumentConverter()
                result = converter.convert(filepath)
                text = result.document.export_to_markdown()
                if text and text.strip():
                    return text
            except Exception as e:
                logger.warning(f"Docling falhou para {filepath}: {e}")

        # Fallback PDF
        if ext == ".pdf":
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(filepath)
                pages = [p.extract_text() or "" for p in reader.pages]
                text = "\n".join(pages)
                if text.strip():
                    return text
            except Exception as e:
                logger.warning(f"PyPDF2 falhou: {e}")

        # TXT, CSV, MD, JSON
        if ext in (".txt", ".csv", ".md", ".log"):
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()

        if ext == ".json":
            with open(filepath, "r", encoding="utf-8") as f:
                data = json.load(f)
                return json.dumps(data, ensure_ascii=False, indent=2)

        # XLSX
        if ext == ".xlsx":
            try:
                import openpyxl
                wb = openpyxl.load_workbook(filepath, read_only=True)
                rows = []
                for ws in wb.worksheets:
                    for row in ws.iter_rows(values_only=True):
                        vals = [str(c) for c in row if c is not None]
                        if vals:
                            rows.append(" | ".join(vals))
                return "\n".join(rows)
            except Exception as e:
                logger.warning(f"openpyxl falhou: {e}")

        # Imagens (OCR)
        if ext in (".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp", ".gif"):
            try:
                from PIL import Image
                import pytesseract
                img = Image.open(filepath)
                text = pytesseract.image_to_string(img, lang="por+eng")
                if text and text.strip():
                    return text
            except Exception as e:
                logger.warning(f"OCR falhou: {e}")

        # DOCX fallback
        if ext == ".docx":
            try:
                import docx
                doc = docx.Document(filepath)
                return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            except Exception as e:
                logger.warning(f"python-docx falhou: {e}")

    except Exception as e:
        logger.error(f"Erro ao extrair texto de {filepath}: {e}")

    return "[Erro: não foi possível ler o arquivo]"


def process_and_store_document(filepath: str, filename: str) -> dict:
    """Processa documento: extrai texto, cria chunks com embeddings e salva no PostgreSQL."""
    db = SessionLocal()
    try:
        ext = os.path.splitext(filename)[1].lower()
        file_size = os.path.getsize(filepath) if os.path.exists(filepath) else 0

        # Criar registro do documento
        doc = Document(
            filename=filename,
            filepath=filepath,
            file_type=ext,
            file_size=file_size,
            status="processing",
        )
        db.add(doc)
        db.flush()

        # Extrair texto
        text = extract_text_from_file(filepath)
        if not text or text.startswith("[Erro"):
            doc.status = "error"
            doc.error_message = "Não foi possível extrair texto"
            db.commit()
            return {"success": False, "error": doc.error_message}

        # Criar chunks (semânticos)
        chunks = split_text_into_chunks(text, chunk_size=400, overlap=50)
        if not chunks:
            doc.status = "error"
            doc.error_message = "Nenhum chunk gerado"
            db.commit()
            return {"success": False, "error": doc.error_message}

        # Gerar embeddings em batch
        embeddings = embed_texts(chunks)
        if embeddings is None:
            logger.warning(f"Embeddings falharam para {filename}, salvando sem vetor")
            embeddings = [None] * len(chunks)

        # Salvar chunks
        for i, (chunk_text, emb) in enumerate(zip(chunks, embeddings)):
            kc = KnowledgeChunk(
                document_id=doc.id,
                chunk_index=i,
                content=chunk_text,
                source=filename,
                category=ext.replace(".", ""),
                embedding=emb,
            )
            db.add(kc)

        doc.num_chunks = len(chunks)
        doc.status = "completed"

        # Status de treinamento
        ts = TrainingStatus(
            document_id=doc.id,
            status="completed",
            message=f"{len(chunks)} chunks processados com embeddings",
        )
        db.add(ts)
        db.commit()

        emb_count = sum(1 for e in embeddings if e is not None)
        logger.info(f"Documento '{filename}': {len(chunks)} chunks, {emb_count} embeddings")

        return {
            "success": True,
            "document_id": doc.id,
            "chunks": len(chunks),
            "embeddings": emb_count,
            "filename": filename,
        }

    except Exception as e:
        db.rollback()
        logger.error(f"Erro ao processar documento {filename}: {e}")
        return {"success": False, "error": str(e)}
    finally:
        db.close()


def get_relevant_context(query: str, max_chunks: int = 6, query_embedding=None) -> str:
    """
    Busca chunks relevantes usando busca vetorial + keywords + nome do documento.
    FILTRA chunks com distância > MAX_COSINE_DISTANCE para não retornar lixo.
    SEMPRE complementa com busca por keywords e source name.
    """
    db = SessionLocal()
    try:
        parts = []
        seen_ids = set()

        # 1. Busca vetorial (cosine distance)
        query_emb = query_embedding if query_embedding is not None else embed_single(query)
        if query_emb is not None:
            distance_col = KnowledgeChunk.embedding.cosine_distance(query_emb)
            results = (
                db.query(KnowledgeChunk, distance_col.label("distance"))
                .filter(KnowledgeChunk.embedding.isnot(None))
                .order_by(distance_col)
                .limit(max_chunks * 3)
                .all()
            )
            if results:
                for r, dist in results:
                    if dist > MAX_COSINE_DISTANCE:
                        continue
                    if r.id in seen_ids:
                        continue
                    src = r.source or "desconhecido"
                    # Injeta metadados de data para a IA saber a recência da informação
                    date_str = r.created_at.strftime("%d/%m/%Y") if r.created_at else "Data desc."
                    parts.append(f"[Fonte: {src} | Data Processamento: {date_str}]\n{r.content}")
                    seen_ids.add(r.id)
                    if len(parts) >= max_chunks:
                        break
                if results:
                    logger.info("KB vetorial: %d chunks (melhor dist=%.3f)", len(parts), results[0][1])

        # 2. Busca por keywords e nome do documento (SEMPRE executa como complemento)
        keywords = [w.lower() for w in query.split() if len(w) >= 2]
        if keywords and len(parts) < max_chunks:
            all_chunks = (
                db.query(KnowledgeChunk)
                .order_by(KnowledgeChunk.created_at.desc())
                .limit(150)
                .all()
            )
            scored = []
            for chunk in all_chunks:
                if chunk.id in seen_ids:
                    continue
                content_lower = chunk.content.lower()
                source_lower = (chunk.source or "").lower()
                # Pontuação por keywords no conteúdo
                score = sum(1 for kw in keywords if kw in content_lower)
                # Bonus: keywords no nome do documento/source (ex: "mikrotik" no filename)
                source_score = sum(2 for kw in keywords if kw in source_lower)
                score += source_score
                if score > 0:
                    scored.append((score, chunk))

            scored.sort(key=lambda x: x[0], reverse=True)
            remaining = max_chunks - len(parts)
            for _, chunk in scored[:remaining]:
                if chunk.id not in seen_ids:
                    src = chunk.source or "desconhecido"
                    # Injeta metadados
                    date_str = chunk.created_at.strftime("%d/%m/%Y") if chunk.created_at else "Data desc."
                    parts.append(f"[Fonte: {src} | Data Processamento: {date_str}]\n{chunk.content}")
                    seen_ids.add(chunk.id)

        # 3. Busca por nome de documento (quando user pede "resumo de X", buscar chunks do doc X)
        if len(parts) < max_chunks:
            import re as _re
            # Detectar padrões como "resumo de X", "sobre o documento X", "conteúdo de X"
            doc_pattern = _re.search(
                r'(?:resumo|resumir|conte[uú]do|sobre|documento|arquivo|pdf|aula)\s+(?:de|do|da|sobre)?\s*(.+)',
                query, _re.IGNORECASE
            )
            if doc_pattern:
                doc_name = doc_pattern.group(1).strip().lower()
                doc_name = _re.sub(r'[?\.]$', '', doc_name).strip()
                if len(doc_name) >= 3:
                    doc_chunks = (
                        db.query(KnowledgeChunk)
                        .filter(KnowledgeChunk.source.isnot(None))
                        .order_by(KnowledgeChunk.chunk_index.asc())
                        .all()
                    )
                    for chunk in doc_chunks:
                        if chunk.id in seen_ids:
                            continue
                        source_lower = (chunk.source or "").lower()
                        # Match pelo nome do documento
                        if doc_name in source_lower or any(
                            w in source_lower for w in doc_name.split() if len(w) >= 3
                        ):
                            src = chunk.source or "desconhecido"
                            # Injeta metadados
                            date_str = chunk.created_at.strftime("%d/%m/%Y") if chunk.created_at else "Data desc."
                            parts.append(f"[Fonte: {src} | Data Processamento: {date_str}]\n{chunk.content}")
                            seen_ids.add(chunk.id)
                            if len(parts) >= max_chunks * 2:
                                break

        if not parts:
            logger.info("KB: nenhum chunk encontrado para '%s'", query[:50])
            return ""

        logger.info("KB total: %d chunks para '%s'", len(parts), query[:50])
        return "\n\n---\n\n".join(parts)

    except Exception as e:
        logger.error(f"Erro na busca de contexto: {e}")
        return ""
    finally:
        db.close()


def get_all_documents() -> List[dict]:
    """Retorna todos os documentos."""
    db = SessionLocal()
    try:
        docs = db.query(Document).order_by(Document.created_at.desc()).all()
        return [
            {
                "id": d.id,
                "uid": d.uid,
                "filename": d.filename,
                "file_type": d.file_type,
                "file_size": d.file_size,
                "num_chunks": d.num_chunks,
                "status": d.status,
                "error_message": d.error_message,
                "created_at": d.created_at.isoformat() if d.created_at else None,
            }
            for d in docs
        ]
    finally:
        db.close()


def get_training_history() -> List[dict]:
    """Retorna histórico de treinamento."""
    db = SessionLocal()
    try:
        statuses = (
            db.query(TrainingStatus)
            .order_by(TrainingStatus.created_at.desc())
            .limit(100)
            .all()
        )
        return [
            {
                "id": s.id,
                "document_id": s.document_id,
                "status": s.status,
                "message": s.message,
                "created_at": s.created_at.isoformat() if s.created_at else None,
            }
            for s in statuses
        ]
    finally:
        db.close()


def delete_document(doc_id: int) -> bool:
    """Remove documento e seus chunks."""
    db = SessionLocal()
    try:
        db.query(KnowledgeChunk).filter(KnowledgeChunk.document_id == doc_id).delete()
        db.query(TrainingStatus).filter(TrainingStatus.document_id == doc_id).delete()
        db.query(Document).filter(Document.id == doc_id).delete()
        db.commit()
        return True
    except Exception as e:
        db.rollback()
        logger.error(f"Erro ao deletar documento {doc_id}: {e}")
        return False
    finally:
        db.close()


def backfill_embeddings(batch_size: int = 32) -> dict:
    """Gera embeddings para chunks que ainda não têm."""
    db = SessionLocal()
    try:
        chunks = (
            db.query(KnowledgeChunk)
            .filter(KnowledgeChunk.embedding.is_(None))
            .limit(batch_size)
            .all()
        )
        if not chunks:
            return {"updated": 0, "message": "Nenhum chunk sem embedding"}

        texts = [c.content[:1000] for c in chunks]
        embeddings = embed_texts(texts)

        if embeddings is None:
            return {"updated": 0, "error": "Não foi possível gerar embeddings"}

        updated = 0
        for chunk, emb in zip(chunks, embeddings):
            if emb is not None:
                chunk.embedding = emb
                updated += 1

        db.commit()
        remaining = (
            db.query(KnowledgeChunk)
            .filter(KnowledgeChunk.embedding.is_(None))
            .count()
        )
        return {"updated": updated, "remaining": remaining}
    except Exception as e:
        db.rollback()
        logger.error(f"Erro no backfill: {e}")
        return {"updated": 0, "error": str(e)}
    finally:
        db.close()